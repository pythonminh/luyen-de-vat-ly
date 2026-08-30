# -*- coding: utf-8 -*-
"""Convert common LaTeX snippets into native editable Word OMML equations."""
from __future__ import annotations
import re
from typing import List, Tuple, Dict, Any
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

GREEK={"alpha":"α","beta":"β","gamma":"γ","delta":"δ","epsilon":"ε","theta":"θ","lambda":"λ","mu":"μ","nu":"ν","xi":"ξ","pi":"π","rho":"ρ","sigma":"σ","tau":"τ","phi":"φ","chi":"χ","psi":"ψ","omega":"ω","Gamma":"Γ","Delta":"Δ","Theta":"Θ","Lambda":"Λ","Xi":"Ξ","Pi":"Π","Sigma":"Σ","Phi":"Φ","Psi":"Ψ","Omega":"Ω"}
SYMBOLS={"times":"×","cdot":"·","pm":"±","le":"≤","leq":"≤","ge":"≥","geq":"≥","neq":"≠","ne":"≠","approx":"≈","infty":"∞","to":"→","rightarrow":"→","leftarrow":"←","Rightarrow":"⇒","Leftarrow":"⇐","leftrightarrow":"↔","Leftrightarrow":"⇔","ldots":"…","dots":"…","cdots":"⋯","in":"∈","notin":"∉","subset":"⊂","subseteq":"⊆","supset":"⊃","supseteq":"⊇","cup":"∪","cap":"∩","emptyset":"∅","forall":"∀","exists":"∃","partial":"∂","nabla":"∇","angle":"∠","perp":"⊥","parallel":"∥"}
ACCENTS={"vec":"→","overrightarrow":"→","widehat":"^","hat":"^","bar":"¯","overline":"¯","underline":"_"}

def _dummy(): return OxmlElement('m:e')
def _mr(parent,text,italic=True,bold=False):
    r=OxmlElement('m:r')
    if not italic or bold:
        rp=OxmlElement('m:rPr')
        if not italic:
            sty=OxmlElement('m:sty'); sty.set(qn('m:val'),'p'); rp.append(sty)
        if bold:
            b=OxmlElement('m:b'); b.set(qn('m:val'),'1'); rp.append(b)
        r.append(rp)
    t=OxmlElement('m:t'); t.text=text; r.append(t); parent.append(r); return r

def _me(parent):
    e=OxmlElement('m:e'); parent.append(e); return e

def _read_group(s,i):
    while i<len(s) and s[i].isspace(): i+=1
    if i>=len(s) or s[i]!='{': return '',i
    depth=0; start=i+1; i+=1
    while i<len(s):
        if s[i]=='\\': i+=2; continue
        if s[i]=='{': depth+=1
        elif s[i]=='}':
            if depth==0: return s[start:i],i+1
            depth-=1
        i+=1
    return s[start:],i

def _plain_math_text(s):
    s=s or ''
    s=re.sub(r'\\(?:textbf|mathbf|mathrm|textit|emph|textrm|text|operatorname|mbox)\s*\{([^{}]*)\}',r'\1',s)
    for k,v in {**GREEK,**SYMBOLS}.items(): s=s.replace('\\'+k,v)
    return s.replace('\\left','').replace('\\right','')

def _atoms(s):
    s=(s or '').strip().replace('\\left','').replace('\\right','')
    out=[]; i=0
    while i<len(s):
        if s[i].isspace():
            d=_dummy(); _mr(d,' ',italic=False); out.append(d[-1]); i+=1; continue
        if s[i]=='{':
            g,i=_read_group(s,i); d=_dummy()
            for n in _atoms(g): d.append(n)
            out.extend(list(d)); continue
        if s[i]=='\\':
            m=re.match(r'\\([A-Za-z]+|.)',s[i:])
            if not m: i+=1; continue
            cmd=m.group(1); i+=len(m.group(0))
            if cmd in ('frac','dfrac','tfrac'):
                num,i=_read_group(s,i); den,i=_read_group(s,i)
                f=OxmlElement('m:f'); f.append(OxmlElement('m:fPr'))
                nu,de=OxmlElement('m:num'),OxmlElement('m:den'); en,ed=_me(nu),_me(de)
                for n in _atoms(num): en.append(n)
                for n in _atoms(den): ed.append(n)
                f.append(nu); f.append(de); out.append(f); continue
            if cmd=='sqrt':
                body,i=_read_group(s,i); rad=OxmlElement('m:rad'); rad.append(OxmlElement('m:radPr')); deg=OxmlElement('m:deg'); deg.append(OxmlElement('m:e')); rad.append(deg); e=_me(rad)
                for n in _atoms(body): e.append(n)
                out.append(rad); continue
            if cmd in ACCENTS:
                body,i=_read_group(s,i); acc=OxmlElement('m:acc'); pr=OxmlElement('m:accPr'); ch=OxmlElement('m:chr'); ch.set(qn('m:val'),ACCENTS[cmd]); pr.append(ch); acc.append(pr); e=_me(acc)
                for n in _atoms(body): e.append(n)
                out.append(acc); continue
            if cmd in ('text','textrm','mathrm','operatorname','mbox'):
                body,i=_read_group(s,i); d=_dummy(); _mr(d,_plain_math_text(body),italic=False); out.append(d[-1]); continue
            if cmd in ('mathbf','boldsymbol'):
                body,i=_read_group(s,i); d=_dummy(); _mr(d,_plain_math_text(body),bold=True); out.append(d[-1]); continue
            val=GREEK.get(cmd) or SYMBOLS.get(cmd)
            if val is not None:
                d=_dummy(); _mr(d,val); out.append(d[-1]); continue
            if cmd in ('{','}','%','_','#','$','&'):
                d=_dummy(); _mr(d,cmd,italic=False); out.append(d[-1]); continue
            d=_dummy(); _mr(d,cmd); out.append(d[-1]); continue
        if s[i] in '^_':
            if not out: i+=1; continue
            op=s[i]; i+=1
            if i<len(s) and s[i]=='{': arg,i=_read_group(s,i)
            else: arg=s[i] if i<len(s) else ''; i+=1
            base=out.pop(); ss=OxmlElement('m:sSup' if op=='^' else 'm:sSub'); be,se=_me(ss),_me(ss); be.append(base)
            for n in _atoms(arg): se.append(n)
            out.append(ss); continue
        d=_dummy(); _mr(d,s[i]); out.append(d[-1]); i+=1
    return out

def add_mixed_latex(paragraph,text,bold_prefix=''):
    if bold_prefix:
        r=paragraph.add_run(bold_prefix); r.bold=True
    text=text or ''
    delim=re.compile(r'(\$\$.*?\$\$|\$.*?\$|\\\(.*?\\\)|\\\[.*?\\\])',re.S)
    pos=0
    for m in delim.finditer(text):
        if m.start()>pos:
            plain=_plain_math_text(text[pos:m.start()])
            if plain: paragraph.add_run(plain)
        token=m.group(0); inner=token[2:-2] if token.startswith('$$') or token.startswith('\\[') or token.startswith('\\(') else token[1:-1]
        om=OxmlElement('m:oMath')
        for node in _atoms(inner): om.append(node)
        paragraph._p.append(om); pos=m.end()
    if pos<len(text):
        tail=_plain_math_text(text[pos:])
        if tail: paragraph.add_run(tail)

def patch_module(ns: Dict[str,Any]):
    """Patch ra_de's Word parser/exporter so LaTeX stays editable as Word equations."""
    def parse(block):
        raw=block or ''; typ='D'
        m=re.search(r'\\choiceTF\b',raw,re.I)
        if m:
            typ='B'; body=raw[:m.start()]; args,_=ns['_extract_n_args'](raw,m.end(),4); options=args[:4]
        else:
            m=re.search(r'\\choice\b',raw,re.I)
            if m:
                typ='A'; body=raw[:m.start()]; args,_=ns['_extract_n_args'](raw,m.end(),4); options=args[:4]
            else:
                m=re.search(r'\\shortans\b',raw,re.I)
                if m:
                    typ='C'; body=raw[:m.start()]; args,_=ns['_extract_n_args'](raw,m.end(),1); options=args[:1]
                else: body=raw; options=[]
        for macro in ('loigiai','dapan'):
            body=ns['_remove_macro_block'](body,macro)
        body=re.sub(r'\\end\s*\{\s*(ex|bt)\s*\}|\\begin\s*\{\s*(ex|bt)\s*\}','',body,flags=re.I)
        body=re.sub(r'^\s*%.*$','',body,flags=re.M).strip()
        options=[re.sub(r'\\True\b','',x,flags=re.I).strip() for x in options]
        return typ,options,body
    ns['_parse_question_block']=parse

    def add_question(doc,index,block,typ):
        qtype,options,body=ns['_parse_question_block'](block)
        if qtype!=typ: qtype=typ
        p=doc.add_paragraph(); p.paragraph_format.space_after=ns['Pt'](4)
        add_mixed_latex(p,body,f'Câu {index}. ')
        if qtype in ('A','B'):
            for j,opt in enumerate(options[:4]):
                p2=doc.add_paragraph(); p2.paragraph_format.left_indent=ns['Cm'](0.7); p2.paragraph_format.space_after=ns['Pt'](2)
                add_mixed_latex(p2,opt,'ABCD'[j]+'. ')
        elif qtype=='C':
            p2=doc.add_paragraph(); p2.paragraph_format.left_indent=ns['Cm'](0.7); r=p2.add_run('Trả lời: '); r.bold=True; p2.add_run('.'*72)
    ns['_add_docx_question']=add_question
PY