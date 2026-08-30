# -*- coding: utf-8 -*-
"""
Ra đề — tạo đề thi từ ngân hàng câu hỏi GitHub (ngan-hang/*.tex), dựa trên
bank_index.json đã có sẵn trong repo (tạo bởi scripts/build_bank_index.py).

Hỗ trợ:
- Chọn A/B/C/D theo từng dạng bài.
- Tạo đề .tex chia 4 phần.
- Tải Word .docx.
- Ghi MA TRẬN ĐỀ vào cả .tex và Word.
- Mở đề .tex đã tạo để chỉnh sửa/bổ sung trực tiếp trên trình duyệt.
"""
from __future__ import annotations

import io
import json
import os
import random
import re
from typing import Any, Dict, List, Optional, Tuple

from flask import Blueprint, Response, request, render_template_string, send_file, redirect
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

bp = Blueprint("ra_de", __name__)
APP_DIR = os.path.dirname(os.path.abspath(__file__))
BANK_INDEX_PATH = os.path.join(APP_DIR, "bank_index.json")

_BLOCK_RE = re.compile(r"\\begin\s*\{\s*(ex|bt)\s*\}(.*?)\\end\s*\{\s*\1\s*\}", re.S | re.I)
_DANGBT_RE = re.compile(r"\\dangbt\s*\{([^{}]*)\}", re.I)
_CHOICE_TF_RE = re.compile(r"\\choiceTF\b", re.I)
_SHORTANS_RE = re.compile(r"\\shortans\b", re.I)
_CHOICE_RE = re.compile(r"\\choice\b", re.I)
LOAI_CAU_LIST = ["Trắc nghiệm", "Đúng sai", "Trả lời ngắn", "Tự luận"]


def _loai_cau_of_block(body: str) -> str:
    body = body or ""
    if _CHOICE_TF_RE.search(body): return "Đúng sai"
    if _SHORTANS_RE.search(body): return "Trả lời ngắn"
    if _CHOICE_RE.search(body): return "Trắc nghiệm"
    return "Tự luận"


def _load_bank_index() -> Dict[str, Any]:
    try:
        with open(BANK_INDEX_PATH, "r", encoding="utf-8") as f: return json.load(f)
    except Exception:
        return {"schema": 1, "total_files": 0, "total_questions": 0, "lessons": []}


def _resolve_tex_path(rel_path: str) -> str:
    rel=(rel_path or "").replace("\\", "/").lstrip("/")
    return os.path.join(APP_DIR, rel.replace("/", os.sep))


def extract_blocks_by_dang(text: str) -> List[Tuple[str,str,str]]:
    text=text or ""; markers=[(m.start(),(m.group(1) or "").strip()) for m in _DANGBT_RE.finditer(text)]
    out=[]; mi=0; current=""
    for m in _BLOCK_RE.finditer(text):
        while mi<len(markers) and markers[mi][0]<m.start(): current=markers[mi][1]; mi+=1
        out.append((current or "Chưa phân dạng", _loai_cau_of_block(m.group(0)), m.group(0)))
    return out


def _read_tex_text(rel_path: str) -> str:
    try:
        with open(_resolve_tex_path(rel_path), "r", encoding="utf-8", errors="replace") as f: return f.read()
    except Exception: return ""


def blocks_grouped_by_dang(rel_path: str) -> Dict[str,List[str]]:
    grouped={}
    for name,_,block in extract_blocks_by_dang(_read_tex_text(rel_path)): grouped.setdefault(name,[]).append(block)
    return grouped


def blocks_grouped_by_dang_loai(rel_path: str) -> Dict[str,Dict[str,List[str]]]:
    grouped={}
    for name,loai,block in extract_blocks_by_dang(_read_tex_text(rel_path)): grouped.setdefault(name,{}).setdefault(loai,[]).append(block)
    return grouped


def dang_loai_counts(rel_path: str) -> Dict[str,Dict[str,int]]:
    grouped=blocks_grouped_by_dang_loai(rel_path)
    return {dang:{lc:len(loai_map.get(lc,[])) for lc in LOAI_CAU_LIST} for dang,loai_map in grouped.items()}


def _build_matrix_rows(selections):
    rows={}
    for path,dang,loai,n in selections:
        rel=(path or "").replace("\\","/")
        lesson=rel.split("ngan-hang/",1)[-1] if "ngan-hang/" in rel else rel
        key=f"{lesson} — {dang}"
        rows.setdefault(key,[0,0,0,0]); rows[key][LOAI_CAU_LIST.index(loai)]+=int(n or 0)
    return [(k,a,b,c,d,a+b+c+d) for k,(a,b,c,d) in rows.items()]


def _matrix_tex(rows):
    if not rows: return ""
    ta=tb=tc=td=tt=0
    lines=["% ==================================================","% MA TRẬN ĐỀ","% ==================================================",r"\begin{center}",r"\textbf{MA TRẬN ĐỀ}",r"\end{center}",r"\renewcommand{\arraystretch}{1.2}",r"\begin{center}",r"\begin{tabular}{|p{7.0cm}|c|c|c|c|c|}",r"\hline",r"\textbf{Bài / Dạng} & \textbf{A} & \textbf{B} & \textbf{C} & \textbf{D} & \textbf{Tổng} \\ \hline"]
    for name,a,b,c,d,t in rows:
        safe=name.replace("&", r"\&")
        lines.append(f"{safe} & {a} & {b} & {c} & {d} & {t} \\\\ \\hline")
        ta+=a; tb+=b; tc+=c; td+=d; tt+=t
    lines.extend([f"\\textbf{{Tổng}} & \\textbf{{{ta}}} & \\textbf{{{tb}}} & \\textbf{{{tc}}} & \\textbf{{{td}}} & \\textbf{{{tt}}} \\\\ \\hline", r"\end{tabular}", r"\end{center}", r"\vspace{0.4cm}", ""])
    return "\n".join(lines)

PAGE_TPL = r'''<!doctype html><html lang="vi"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>📝 Ra đề</title><style>body{font-family:Arial;background:#f6f8fa;color:#172b4d;margin:0}.wrap{max-width:1200px;margin:20px auto;padding:0 16px 80px}.card{background:#fff;border:1px solid #d0d7de;border-radius:12px;padding:16px;margin-bottom:14px}.lesson{border:1px solid #e2e8f0;border-radius:10px;padding:10px 12px;margin:8px 0}.lesson summary{cursor:pointer;font-weight:700}.dang-block{border-top:1px dashed #e5e7eb;padding:9px 4px}.type-grid{display:grid;grid-template-columns:repeat(4,minmax(180px,1fr));gap:8px}.type-box{display:flex;align-items:center;gap:7px;border:1px solid #dbe3ec;border-radius:8px;padding:8px}.type-box input[type=number]{width:54px;margin-left:auto;padding:5px}.muted{color:#64748b;font-size:13px}.badge{background:#eef2ff;color:#3730a3;border-radius:999px;padding:2px 8px;font-size:12px;margin-left:6px}.lesson-tools{margin:8px 0}button,.btn{background:#1976d2;color:#fff;border:0;padding:9px 14px;border-radius:8px;font-weight:700;cursor:pointer;text-decoration:none;display:inline-block}button.small{padding:5px 9px;font-size:12px;background:#475569;margin-right:4px}.btn.secondary{background:#64748b}.search{width:100%;box-sizing:border-box;padding:8px 10px;border:1px solid #cbd5e1;border-radius:8px;margin-bottom:10px}@media(max-width:850px){.type-grid{grid-template-columns:repeat(2,minmax(170px,1fr))}}@media(max-width:520px){.type-grid{grid-template-columns:1fr}}</style></head><body><div class="wrap"><div class="card"><h1>📝 Ra đề từ ngân hàng GitHub</h1><p class="muted">Nguồn: <b>{{ total_files }}</b> file · <b>{{ total_questions }}</b> câu hỏi. Tick A/B/C/D ở từng dạng bài rồi nhập số câu.</p></div><form method="post" action="{{ url_for('ra_de.generate') }}"><div class="card"><input class="search" id="q" placeholder="🔎 Lọc..." onkeyup="filterLessons()"><label class="muted">Tên đề: <input type="text" name="ten_de" placeholder="Đề ôn tập..."></label></div>{% for group_key, lessons in groups.items() %}<div class="card lesson-group" data-key="{{ group_key|lower }}"><h3>{{ group_key }}</h3>{% for item in lessons %}<details class="lesson"><summary>{{ item.BaiHoc or item.path }} <span class="badge">{{ item.questions }} câu</span></summary><div class="lesson-tools"><button type="button" class="small" onclick="setLessonType(this,'A',true)">☑ Chọn A</button><button type="button" class="small" onclick="setLessonType(this,'B',true)">☑ Chọn B</button><button type="button" class="small" onclick="setLessonType(this,'C',true)">☑ Chọn C</button><button type="button" class="small" onclick="setLessonType(this,'D',true)">☑ Chọn D</button><button type="button" class="small" onclick="setLessonType(this,'ALL',true)">☑ Chọn cả 4</button><button type="button" class="small" onclick="setLessonType(this,'ALL',false)">☐ Bỏ chọn</button></div>{% for dang, counts in item.dang_loai.items() %}{% set dang_idx=loop.index0 %}<div class="dang-block"><div><b>{{ dang }}</b> <span class="muted">— A: {{ counts.get("Trắc nghiệm",0) }}, B: {{ counts.get("Đúng sai",0) }}, C: {{ counts.get("Trả lời ngắn",0) }}, D: {{ counts.get("Tự luận",0) }}</span></div><div class="type-grid">{% for loai_cau,code in [("Trắc nghiệm","A"),("Đúng sai","B"),("Trả lời ngắn","C"),("Tự luận","D")] %}{% set n2=counts.get(loai_cau,0) %}<label class="type-box"><input type="checkbox" class="type-check" data-type="{{ code }}" data-max="{{ n2 }}" data-key="{{ item._ra_index }}|{{ dang_idx }}|{{ code }}" {% if n2<=0 %}disabled{% endif %} onchange="toggleCount(this)"><span><b>{{ code }}. {{ loai_cau }}</b> <span class="muted">({{ n2 }} câu)</span></span><input type="number" min="0" max="{{ n2 }}" value="0" class="count-input" disabled></label>{% endfor %}</div></div>{% endfor %}</details>{% endfor %}</div>{% endfor %}<input type="hidden" name="selections" id="selections"><div class="card"><button type="submit">✅ Tạo đề (.tex)</button><a class="btn secondary" href="{{ url_for('ra_de.home') }}">↺ Làm lại</a></div></form></div><script>function toggleCount(c){let i=c.closest('.type-box').querySelector('.count-input');i.disabled=!c.checked;if(c.checked){if(+i.value<=0)i.value=1;let m=+c.dataset.max;if(+i.value>m)i.value=m}else i.value=0}function setLessonType(b,t,x){let l=b.closest('.lesson');l.querySelectorAll('.type-check').forEach(c=>{if((t==='ALL'||c.dataset.type===t)&&!c.disabled){c.checked=x;toggleCount(c)}})}function filterLessons(){let q=(document.getElementById('q').value||'').toLowerCase();document.querySelectorAll('.lesson-group').forEach(g=>g.style.display=!q||g.innerText.toLowerCase().includes(q)?'':'none')}document.querySelector('form').addEventListener('submit',function(){let s=[];document.querySelectorAll('.type-check:checked').forEach(c=>{let i=c.closest('.type-box').querySelector('.count-input'),n=+i.value;if(n>0)s.push({k:c.dataset.key,n:n})});document.getElementById('selections').value=JSON.stringify(s)})</script></body></html>'''

RESULT_TPL = r'''<!doctype html><html lang="vi"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>📝 Đề đã tạo</title><style>body{font-family:Arial;background:#f6f8fa;color:#172b4d}.wrap{max-width:1000px;margin:20px auto;padding:0 16px 60px}.card{background:#fff;border:1px solid #d0d7de;border-radius:12px;padding:16px;margin-bottom:14px}pre{white-space:pre-wrap;word-break:break-word;background:#0f172a;color:#e2e8f0;padding:14px;border-radius:10px;max-height:70vh;overflow:auto}.btn{background:#1976d2;color:#fff;text-decoration:none;padding:10px 18px;border-radius:8px;font-weight:700;display:inline-block;margin-right:7px}.secondary{background:#64748b}</style></head><body><div class="wrap"><div class="card"><h2>✅ Đã tạo đề — {{ total }} câu</h2><p><a class="btn" href="{{ url_for('ra_de.download',token=token) }}">⬇️ Tải file .tex</a><a class="btn" href="{{ url_for('ra_de.download_word',token=token) }}">📝 Tải Word (.docx)</a><a class="btn" href="{{ url_for('ra_de.edit',token=token) }}">✏️ Chỉnh sửa .tex</a><a class="btn secondary" href="{{ url_for('ra_de.home') }}">↺ Tạo đề khác</a></p></div><div class="card"><h3>Xem trước</h3><pre>{{ content }}</pre></div></div></body></html>'''

_LAST_GENERATED={}

def home():
    data=_load_bank_index(); lessons=data.get('lessons') or []; groups={}
    for lesson_idx,item0 in enumerate(lessons):
        if not isinstance(item0,dict): continue
        item=dict(item0); item['_ra_index']=lesson_idx; item['dang_loai']=dang_loai_counts(item.get('path') or '')
        key=f"{item.get('Mon') or 'Khác'} · Lớp {item.get('Lop') or ''} · {item.get('Chuong') or ''}".strip()
        groups.setdefault(key,[]).append(item)
    return render_template_string(PAGE_TPL,groups=groups,total_files=data.get('total_files',len(lessons)),total_questions=data.get('total_questions',0))
bp.add_url_rule('/ra-de','home',home,methods=['GET'])

def _parse_compact_selections(raw,lessons):
    try: data=json.loads(raw or '[]')
    except Exception: return []
    out=[]; names={'A':'Trắc nghiệm','B':'Đúng sai','C':'Trả lời ngắn','D':'Tự luận'}
    for row in data:
        if not isinstance(row,dict): continue
        try: key=str(row.get('k') or ''); n=int(row.get('n') or 0); bits=key.split('|'); li=int(bits[0]); di=int(bits[1]); code=bits[2].upper()
        except Exception: continue
        if code not in names or n<=0 or not(0<=li<len(lessons)): continue
        path=str(lessons[li].get('path') or ''); dm=dang_loai_counts(path); dns=list(dm.keys())
        if not path or not(0<=di<len(dns)): continue
        dang=dns[di]; loai=names[code]; n=min(n,int(dm[dang].get(loai,0)))
        if n>0: out.append((path,dang,loai,n))
    return out

@bp.route('/ra-de/generate',methods=['POST'])
def generate():
    ten_de=(request.form.get('ten_de') or 'Đề ôn tập').strip(); lessons=_load_bank_index().get('lessons') or []
    selections=_parse_compact_selections(request.form.get('selections') or '[]',lessons); wanted={x:{} for x in LOAI_CAU_LIST}
    for path,dang,loai,n in selections: wanted[loai].setdefault(path,{})[dang]=n
    picked={x:[] for x in LOAI_CAU_LIST}
    for loai in LOAI_CAU_LIST:
        for path,dangs in wanted[loai].items():
            g=blocks_grouped_by_dang_loai(path)
            for dang,n in dangs.items():
                pool=(g.get(dang) or {}).get(loai) or []; picked[loai].extend(random.sample(pool,min(n,len(pool)))) if pool else None
    for x in picked: random.shuffle(picked[x])
    counts={x:len(v) for x,v in picked.items()}; total=sum(counts.values()); titles={'Trắc nghiệm':'PHẦN A. TRẮC NGHIỆM 4 LỰA CHỌN','Đúng sai':'PHẦN B. TRẮC NGHIỆM ĐÚNG / SAI','Trả lời ngắn':'PHẦN C. TRẢ LỜI NGẮN','Tự luận':'PHẦN D. TỰ LUẬN'}; codes={'Trắc nghiệm':'A','Đúng sai':'B','Trả lời ngắn':'C','Tự luận':'D'}
    matrix=_matrix_tex(_build_matrix_rows(selections)); parts=[]
    for loai in LOAI_CAU_LIST:
        if not picked[loai]: continue
        parts += ["% ==================================================\n",f"% {titles[loai]}\n","% ==================================================\n\n"]
        for i,b in enumerate(picked[loai],1): parts.append(f"% ===== {codes[loai]} - Câu {i} =====\n{b.strip()}\n\n")
    content=(f"% ===== {ten_de} =====\n% Tự động tạo bởi Ra đề — {total} câu\n% A={counts['Trắc nghiệm']} | B={counts['Đúng sai']} | C={counts['Trả lời ngắn']} | D={counts['Tự luận']}\n\n"+matrix+''.join(parts))
    if total==0: content+="% Chưa có câu nào được chọn.\n% Hãy tick A/B/C/D và nhập số câu > 0.\n"
    token=f"{random.randint(100000,999999)}"; _LAST_GENERATED[token]=content
    return render_template_string(RESULT_TPL,content=content,total=total,token=token)


def _read_braced_arg(s,start):
    i=start
    while i<len(s) and s[i].isspace(): i+=1
    if i>=len(s) or s[i]!='{': return '',start
    depth=0; begin=i+1; i+=1
    while i<len(s):
        if s[i]=='\\': i+=2; continue
        if s[i]=='{': depth+=1
        elif s[i]=='}':
            if depth==0: return s[begin:i],i+1
            depth-=1
        i+=1
    return s[begin:],i


def _extract_n_args(s,start,count):
    out=[]; pos=start
    for _ in range(count):
        a,pos2=_read_braced_arg(s,pos)
        if pos2==pos: break
        out.append(a); pos=pos2
    return out


def _remove_macro_block(s,macro):
    pat=re.compile(r'\\'+re.escape(macro)+r'\s*\{',re.I)
    while True:
        m=pat.search(s)
        if not m: return s
        _,end=_read_braced_arg(s,m.end()-1); s=s[:m.start()]+s[end:]


def _tex_to_word_text(s):
    s=s or ''; s=re.sub(r'%.*$','',s,flags=re.M); s=re.sub(r'\\begin\s*\{\s*tikzpicture\s*\}.*?\\end\s*\{\s*tikzpicture\s*\}','[Hình minh họa]',s,flags=re.S|re.I)
    s=re.sub(r'\\(?:textbf|mathbf|mathrm|textit|emph|textrm|text|underline)\s*\{([^{}]*)\}',r'\1',s)
    for a,b in [(r'\quad','    '),(r'\qquad','        '),(r'\,',' '),(r'\;',' '),(r'\!',''),(r'\times','×'),(r'\cdot','·'),(r'\le','≤'),(r'\ge','≥'),(r'\neq','≠'),(r'\approx','≈'),(r'\pm','±'),(r'\to','→'),(r'\rightarrow','→'),(r'\Rightarrow','⇒'),(r'\infty','∞'),(r'\alpha','α'),(r'\beta','β'),(r'\gamma','γ'),(r'\Delta','Δ'),(r'\Omega','Ω')]: s=s.replace(a,b)
    for _ in range(6):
        old=s; s=re.sub(r'\\(?:d?frac)\s*\{([^{}]*)\}\s*\{([^{}]*)\}',r'(\1)/(\2)',s)
        if old==s: break
    s=s.replace('$$',''); s=re.sub(r'\$(.*?)\$',r'\1',s,flags=re.S); s=re.sub(r'\\(?:left|right|big|Big|bigg|Bigg)\b','',s); s=re.sub(r'\\[a-zA-Z]+\*?(?!\w)','',s)
    return re.sub(r'\n{3,}','\n\n',s).strip()


def _parse_question_block(block):
    raw=block or ''; typ='D'; opts=[]; m=re.search(r'\\choiceTF\b',raw,re.I)
    if m: typ='B'; body=raw[:m.start()]; opts=_extract_n_args(raw,m.end(),4)
    else:
        m=re.search(r'\\choice\b',raw,re.I)
        if m: typ='A'; body=raw[:m.start()]; opts=_extract_n_args(raw,m.end(),4)
        else:
            m=re.search(r'\\shortans\b',raw,re.I)
            if m: typ='C'; body=raw[:m.start()]; opts=_extract_n_args(raw,m.end(),1)
            else: body=raw
    for macro in ('loigiai','dapan'): body=_remove_macro_block(body,macro)
    body=re.sub(r'\\(?:begin|end)\s*\{\s*(?:ex|bt)\s*\}','',body,flags=re.I); body=re.sub(r'^\s*%.*$','',body,flags=re.M).strip()
    return typ,[_tex_to_word_text(re.sub(r'\\True\b','',x,flags=re.I)) for x in opts],_tex_to_word_text(body)


def _blocks_from_generated(content):
    r={'A':[],'B':[],'C':[],'D':[]}
    for m in _BLOCK_RE.finditer(content or ''):
        typ,_,_=_parse_question_block(m.group(0)); r[typ].append(m.group(0))
    return r


def _build_word_file(content,ten_de):
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Cm(1.6); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8); doc.styles['Normal'].font.name='Arial'; doc.styles['Normal'].font.size=Pt(11)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(ten_de or 'Đề ôn tập'); r.bold=True; r.font.size=Pt(16)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=p.add_run('Đề được tạo tự động từ ngân hàng câu hỏi GitHub'); rr.italic=True
    grouped=_blocks_from_generated(content)
    for code,heading in [('A','PHẦN A. TRẮC NGHIỆM 4 LỰA CHỌN'),('B','PHẦN B. TRẮC NGHIỆM ĐÚNG / SAI'),('C','PHẦN C. TRẢ LỜI NGẮN'),('D','PHẦN D. TỰ LUẬN')]:
        if not grouped.get(code): continue
        h=doc.add_paragraph(); r=h.add_run(heading); r.bold=True; r.font.size=Pt(13)
        for i,b in enumerate(grouped[code],1):
            _,opts,body=_parse_question_block(b); p=doc.add_paragraph(); p.add_run(f'Câu {i}. ').bold=True; p.add_run(body)
            if code in ('A','B'):
                for j,opt in enumerate(opts[:4]):
                    q=doc.add_paragraph(); q.paragraph_format.left_indent=Cm(.7); q.add_run(f"{'ABCD'[j]}. ").bold=True; q.add_run(opt)
            elif code=='C':
                q=doc.add_paragraph(); q.paragraph_format.left_indent=Cm(.7); q.add_run('Trả lời: ').bold=True; q.add_run('.'*72)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

@bp.route('/ra-de/download-word/<token>')
def download_word(token):
    content=_LAST_GENERATED.get(token,'')
    if not content:return Response('Không tìm thấy đề đã tạo hoặc đề đã hết phiên.',404)
    m=re.search(r'^%\s*=====\s*(.*?)\s*=====',content,flags=re.M); ten_de=m.group(1).strip() if m else 'Đề ôn tập'
    try: buf=_build_word_file(content,ten_de)
    except Exception as exc: return Response(f'Lỗi tạo file Word: {type(exc).__name__}: {exc}',500,mimetype='text/plain; charset=utf-8')
    safe=re.sub(r'[\\/:*?"<>|]+','_',ten_de).strip() or 'de_thi'; return send_file(buf,mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',as_attachment=True,download_name=f'{safe}.docx')


@bp.route('/ra-de/download/<token>')
def download(token):
    buf=io.BytesIO(_LAST_GENERATED.get(token,'').encode('utf-8')); buf.seek(0); return send_file(buf,mimetype='text/plain; charset=utf-8',as_attachment=True,download_name='de_thi.tex')


@bp.route('/ra-de/edit/<token>',methods=['GET'])
def edit(token):
    content=_LAST_GENERATED.get(token,'')
    if not content:return Response('Không tìm thấy đề hoặc phiên chỉnh sửa đã hết.',404)
    tpl=r'''<!doctype html><html lang="vi"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>✏️ Chỉnh sửa đề</title><style>body{font-family:Arial;background:#f6f8fa;color:#172b4d}.wrap{max-width:1200px;margin:20px auto;padding:0 16px 70px}.card{background:#fff;border:1px solid #d0d7de;border-radius:12px;padding:16px}textarea{width:100%;min-height:72vh;box-sizing:border-box;border:1px solid #94a3b8;border-radius:10px;padding:14px;font:14px/1.45 Consolas,monospace;background:#0f172a;color:#e2e8f0;resize:vertical}.btn,button{border:0;border-radius:8px;padding:10px 15px;background:#1976d2;color:#fff;font-weight:700;text-decoration:none;display:inline-block;cursor:pointer;margin:10px 8px 0 0}.green{background:#15803d}.gray{background:#64748b}.ok{padding:9px 12px;background:#ecfdf5;border:1px solid #86efac;color:#166534;border-radius:8px;margin:10px 0}</style></head><body><div class="wrap"><div class="card"><h2>✏️ Chỉnh sửa / bổ sung câu hỏi</h2><p>Sửa mã LaTeX, thêm câu hỏi, đổi thứ tự hoặc chỉnh ma trận. Sau đó bấm <b>Lưu thay đổi</b>.</p>{% if saved %}<div class="ok">✅ Đã lưu thay đổi.</div>{% endif %}<form method="post" action="{{ url_for('ra_de.save_edit',token=token) }}"><textarea name="content" spellcheck="false">{{ content }}</textarea><div><button type="submit">💾 Lưu thay đổi</button><a class="btn" href="{{ url_for('ra_de.download_edit',token=token) }}">⬇️ Tải .tex</a><a class="btn green" href="{{ url_for('ra_de.download_word',token=token) }}">📝 Tải Word</a><a class="btn gray" href="{{ url_for('ra_de.home') }}">↺ Tạo đề khác</a></div></form></div></div></body></html>'''
    return render_template_string(tpl,token=token,content=content,saved=(request.args.get('saved')=='1'))


@bp.route('/ra-de/save-edit/<token>',methods=['POST'])
def save_edit(token):
    if token not in _LAST_GENERATED:return Response('Không tìm thấy đề hoặc phiên chỉnh sửa đã hết.',404)
    content=request.form.get('content','')
    if not content.strip():return Response('Nội dung TeX không được để trống.',400)
    _LAST_GENERATED[token]=content
    return redirect(f'/ra-de/edit/{token}?saved=1')


@bp.route('/ra-de/edit/<token>/download')
def download_edit(token):
    content=_LAST_GENERATED.get(token,'')
    if not content:return Response('Không tìm thấy đề.',404)
    buf=io.BytesIO(content.encode('utf-8')); buf.seek(0)
    return send_file(buf,mimetype='text/plain; charset=utf-8',as_attachment=True,download_name='de_thi_chinh_sua.tex')
